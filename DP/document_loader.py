"""
文档加载层：按格式做针对性解析，统一输出「Markdown 标题结构」的 Document，
供 semantic_splitter 做结构感知切分。

- json：问答数组逐条展开（instruction/input/output -> 问题/答案），保持原子性
- txt/md：直接读取；纯文本无 Markdown 标题时按启发式规则提升标题
- pdf：逐页抽取（优先 pdfplumber，含表格转 Markdown；退回 pypdf）->
       去页眉页脚/页码 -> 合并整篇并注入 [[PAGE=n]] 页码标记 -> 启发式标题提升
- docx：优先 python-docx（按 Heading 样式还原标题层级、表格转 Markdown）；
        退回 docx2txt + 启发式标题提升
- csv：CSVLoader 按行展开（每行带表头字段名）
"""

import hashlib
import json
import re
from collections import Counter
from pathlib import Path
from typing import Dict, List, Optional, Tuple

from langchain_community.document_loaders import CSVLoader, TextLoader
from langchain_core.documents import Document

from RAG.config.logger_runtime import get_logger

logger = get_logger("rag")

# 仓库根目录（含 RAG 包与 assets）：source / doc_hash.json 键统一为相对此根的 POSIX 路径，如 RAG/assets/data/xxx.json
_REPO_ROOT = Path(__file__).resolve().parents[2]

# 支持的文件后缀（pdf/docx/json 走专用加载逻辑）
SUPPORTED_SUFFIXES = {".txt", ".md", ".pdf", ".docx", ".csv", ".json"}


def canonical_source_path(file_path: Path) -> str:
    """
    将数据文件路径规范为相对仓库根的写法（RAG/assets/data/...），与 doc_hash.json 键一致。
    若文件不在仓库根之下则退回绝对路径（POSIX 字符串）。
    """
    abs_p = file_path.resolve()
    try:
        rel = abs_p.relative_to(_REPO_ROOT)
    except ValueError:
        logger.warning("路径不在仓库根 %s 之下，source 使用绝对路径: %s", _REPO_ROOT, abs_p)
        return str(abs_p).replace("\\", "/")
    return str(rel).replace("\\", "/")


def resolve_doc_tag(file_path: Path, data_root: Path, default_tag: str = "general") -> str:
    """
    根据文件相对 data_dir 的一级子目录推断标签：
      - data/ls6/LS6.txt          -> "ls6"
      - data/ls6/sub/x.txt        -> "ls6"（仅取一级子目录）
      - data/car_info.json        -> default_tag（直接位于 data_dir 下）
    无法定位到 data_root 之下时退回 default_tag。
    """
    try:
        rel = file_path.resolve().relative_to(data_root.resolve())
    except ValueError:
        return default_tag
    parts = rel.parts
    if len(parts) > 1:
        return parts[0]
    return default_tag


def normalize_source_key(key: str) -> str:
    """将 doc_hash.json 等处的 source 键规范为与 canonical_source_path 相同的相对路径（兼容旧版绝对路径）。"""
    if not key:
        return key
    key = key.strip().replace("\\", "/")
    p = Path(key)
    if p.is_absolute():
        try:
            return str(p.resolve().relative_to(_REPO_ROOT)).replace("\\", "/")
        except ValueError:
            return str(p.resolve()).replace("\\", "/")
    return key.lstrip("./")


def alternate_source_keys(key: str) -> set[str]:
    """
    返回与同一逻辑文件等价的 source 字符串集合（相对仓库根、绝对路径），
    用于删除 sqlite 中旧绝对路径行或对齐历史索引。
    """
    if not key:
        return set()
    raw = str(key).strip().replace("\\", "/")
    out: set[str] = {raw, normalize_source_key(raw)}
    p = Path(raw)
    if p.is_absolute():
        try:
            rel = str(p.resolve().relative_to(_REPO_ROOT)).replace("\\", "/")
            out.add(rel)
        except ValueError:
            pass
    else:
        try:
            abs_p = str((_REPO_ROOT / raw).resolve()).replace("\\", "/")
            out.add(abs_p)
        except (ValueError, OSError):
            pass
    return {k for k in out if k}


def compute_file_hash(file_path: Path) -> str:
    """计算文件哈希，用于增量模式下判定文件是否变化。"""
    sha = hashlib.sha256()
    with file_path.open("rb") as file_obj:
        for block in iter(lambda: file_obj.read(1024 * 1024), b""):
            sha.update(block)
    return sha.hexdigest()


def load_hash_registry(index_dir: Optional[str]) -> Dict[str, str]:
    """
    从 index_store/doc_hash.json 读取 source -> doc_hash，用于增量模式下判断是否需重新加载文件。
    键统一为相对仓库根的路径（如 RAG/assets/data/...）；若文件不存在或解析失败则返回空字典。
    """
    if not index_dir:
        return {}

    doc_hash_path = Path(index_dir) / "doc_hash.json"
    if not doc_hash_path.exists():
        return {}
    try:
        payload = json.loads(doc_hash_path.read_text(encoding="utf-8"))
    except Exception as exc:
        logger.warning("读取 doc_hash.json 失败，将忽略增量缓存: %s, err=%s", doc_hash_path, exc)
        return {}
    if not isinstance(payload, dict):
        return {}
    return {
        normalize_source_key(str(k)): str(v)
        for k, v in payload.items()
        if k and v
    }


# =============================================================================
# 文本清洗
# =============================================================================

def _compact_fragmented_line(line: str) -> str:
    """修复 PDF 抽取后被拆散的字符序列，如 '4 0 1 . 0'。"""
    tokens = line.split()
    if len(tokens) < 4:
        return line
    single_char_ratio = sum(1 for token in tokens if len(token) == 1) / len(tokens)
    if single_char_ratio < 0.6:
        return line
    return "".join(tokens)


def clean_text(text: str) -> str:
    """
    清洗 PDF/文本中的控制字符、空字节与明显噪音行。
    保留：Markdown 标题、表格行（| 开头）、段落空行（连续空行压成一个）。
    """
    if not text:
        return ""

    text = text.replace("\x00", "")
    text = text.replace("\ufeff", "").replace("\u200b", "")
    text = re.sub(r"[\x01-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]", " ", text)
    text = text.replace("\r\n", "\n").replace("\r", "\n")

    cleaned_lines: List[str] = []
    for raw_line in text.split("\n"):
        stripped = raw_line.strip()
        if not stripped:
            # 保留段落边界（连续空行后续统一压缩）
            cleaned_lines.append("")
            continue
        # 表格行保持原样（分隔行 |---| 无「有效字符」但结构必需）
        if stripped.startswith("|"):
            cleaned_lines.append(stripped)
            continue
        line = re.sub(r"\s+", " ", stripped)
        line = _compact_fragmented_line(line)
        visible = re.sub(r"\s+", "", line)
        meaningful_chars = re.findall(r"[\u4e00-\u9fffA-Za-z0-9]", visible)
        if not meaningful_chars and len(visible) >= 3:
            continue
        if meaningful_chars and (len(meaningful_chars) / max(len(visible), 1)) < 0.25:
            continue
        cleaned_lines.append(line)

    text = "\n".join(cleaned_lines)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# =============================================================================
# 启发式标题识别（用于 PDF / 无结构纯文本）
# =============================================================================

# (正则, 标题级别)：章节编号类
_NUMBERED_HEADING_PATTERNS: List[Tuple[re.Pattern, int]] = [
    (re.compile(r"^第[一二三四五六七八九十百0-9]+[章篇部](?:\s|$)"), 1),
    (re.compile(r"^第[一二三四五六七八九十百0-9]+[节条](?:\s|$)"), 2),
    (re.compile(r"^[一二三四五六七八九十]+\s*[、.．]"), 2),
    (re.compile(r"^[（(][一二三四五六七八九十]+[）)]"), 3),
    (re.compile(r"^\d+\.\d+\.\d+(?:[\s、.．]|$)"), 4),
    (re.compile(r"^\d+\.\d+(?:[\s、.．]|$)"), 3),
]

_BULLET_PREFIX_RE = re.compile(r"^\s*([-*•·]|\d+[.、）)])\s+")
_TERMINAL_PUNCT = "。！？，；、!?,;…"


def _looks_like_markdown(text: str) -> bool:
    """已有 >=2 个 Markdown 标题行时认为文档自带结构，不再做启发式提升。"""
    return len(re.findall(r"^#{1,6}\s+\S", text, flags=re.MULTILINE)) >= 2


_PAGE_MARKER_LINE_RE = re.compile(r"^\[\[PAGE=\d+\]\]\s*$")
# 句末边界符：上一行以这些字符结尾时，当前短行可视为新章节标题
_BOUNDARY_END_CHARS = "。！？!?…；;："
# 数据行结尾（数字/单位收尾，如 "450km" "22.99万"）：也视为段落边界
_DATA_LINE_END_RE = re.compile(r"\d+(?:\.\d+)?\s*[a-zA-Z%‰°℃入万亿元件个条款]{0,6}$")


def _heading_level_for_line(line: str, prev_boundary: bool, next_line: str) -> Optional[int]:
    """
    判断一行是否应提升为标题，返回标题级别；不满足返回 None。
    规则：
      1. 章节编号（第X章 / 一、 / 1.1 等）且行长 <= 40、无句末标点结尾；
      2. 短行标题：处于段落边界（前一行为空/句末/标题/页码标记）、行长 <= 20、
         不含句中标点、词数 <= 2（排除表格式行）、非列表项、下一行有正文。
    """
    stripped = line.strip()
    if not stripped or len(stripped) > 40:
        return None
    if stripped.startswith("#") or stripped.startswith("|"):
        return None
    if _BULLET_PREFIX_RE.match(stripped):
        return None
    if stripped[-1] in _TERMINAL_PUNCT:
        return None

    for pattern, level in _NUMBERED_HEADING_PATTERNS:
        if pattern.match(stripped):
            return level

    # 短行标题：如「产品定位」「超级增程特点」
    if (
        prev_boundary
        and len(stripped) <= 20
        and len(stripped.split()) <= 2
        and next_line.strip()
        and not re.search(r"[。，；：,;:]", stripped)
        and not re.fullmatch(r"[\d\s.%-]+", stripped)
    ):
        return 2
    return None


def promote_headings(text: str) -> str:
    """将无 Markdown 结构文本中的章节编号行/短标题行提升为 Markdown 标题。"""
    if not text or _looks_like_markdown(text):
        return text
    lines = text.split("\n")
    out: List[str] = []
    prev_boundary = True  # 文档开头视为边界
    for idx, line in enumerate(lines):
        next_line = lines[idx + 1] if idx + 1 < len(lines) else ""
        # 页码标记行原样保留，且视为段落边界
        if _PAGE_MARKER_LINE_RE.match(line):
            out.append(line)
            prev_boundary = True
            continue
        stripped = line.strip()
        if not stripped:
            out.append(line)
            prev_boundary = True
            continue
        level = _heading_level_for_line(line, prev_boundary, next_line)
        if level:
            out.append(f"{'#' * level} {stripped}")
            prev_boundary = True
        else:
            out.append(line)
            prev_boundary = (
                stripped[-1] in _BOUNDARY_END_CHARS
                or bool(_DATA_LINE_END_RE.search(stripped))
            )
    return "\n".join(out)


# =============================================================================
# PDF 加载
# =============================================================================

def _render_table_markdown(rows: List[List[Optional[str]]]) -> str:
    """将 pdfplumber 表格转为 Markdown 表格。"""
    if not rows:
        return ""
    norm_rows = [
        [re.sub(r"\s+", " ", (cell or "")).strip() for cell in row] for row in rows
    ]
    norm_rows = [row for row in norm_rows if any(row)]
    if not norm_rows:
        return ""
    width = max(len(r) for r in norm_rows)
    norm_rows = [r + [""] * (width - len(r)) for r in norm_rows]
    lines = ["| " + " | ".join(norm_rows[0]) + " |"]
    lines.append("|" + "---|" * width)
    for row in norm_rows[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


def _extract_pdf_pages(file_path: Path) -> List[str]:
    """
    逐页抽取 PDF 文本：优先 pdfplumber（表格区域单独转 Markdown、正文排除表格避免重复），
    未安装或失败时退回 pypdf。
    """
    try:
        import pdfplumber  # type: ignore
    except ImportError:
        pdfplumber = None

    if pdfplumber is not None:
        try:
            pages: List[str] = []
            with pdfplumber.open(str(file_path)) as pdf:
                for page in pdf.pages:
                    tables = page.find_tables()
                    if tables:
                        bboxes = [t.bbox for t in tables]

                        def _outside_tables(obj, _bboxes=bboxes):
                            v_mid = (obj["top"] + obj["bottom"]) / 2
                            h_mid = (obj["x0"] + obj["x1"]) / 2
                            for x0, top, x1, bottom in _bboxes:
                                if x0 <= h_mid < x1 and top <= v_mid < bottom:
                                    return False
                            return True

                        body = page.filter(_outside_tables).extract_text() or ""
                        table_md = "\n\n".join(
                            md
                            for t in tables
                            if (md := _render_table_markdown(t.extract()))
                        )
                        pages.append(f"{body}\n\n{table_md}".strip())
                    else:
                        pages.append(page.extract_text() or "")
            return pages
        except Exception as exc:  # noqa: BLE001
            logger.warning("pdfplumber 解析失败，退回 pypdf: %s, err=%s", file_path, exc)

    from pypdf import PdfReader

    reader = PdfReader(str(file_path))
    return [(page.extract_text() or "") for page in reader.pages]


_PAGE_NUMBER_LINE_RE = re.compile(
    r"^\s*(?:[-–—]\s*)?(?:第\s*)?\d{1,4}(?:\s*页)?(?:\s*/\s*\d{1,4})?(?:\s*[-–—])?\s*$"
)


def _strip_repeated_page_lines(pages: List[str]) -> List[str]:
    """
    去除页眉/页脚：统计各页首尾 3 行，出现率 >= 60%（至少 2 页）的行（数字归一后）全部删除；
    另删除纯页码行。
    """
    if len(pages) < 2:
        return [
            "\n".join(
                ln for ln in p.split("\n") if not _PAGE_NUMBER_LINE_RE.match(ln)
            )
            for p in pages
        ]

    def _norm(line: str) -> str:
        return re.sub(r"\d+", "#", re.sub(r"\s+", " ", line.strip()))

    counter: Counter = Counter()
    for page in pages:
        lines = [ln for ln in page.split("\n") if ln.strip()]
        edge = lines[:3] + lines[-3:]
        for ln in {_norm(x) for x in edge if x.strip()}:
            counter[ln] += 1

    threshold = max(2, int(len(pages) * 0.6))
    repeated = {ln for ln, cnt in counter.items() if cnt >= threshold and ln}

    result: List[str] = []
    for page in pages:
        kept: List[str] = []
        lines = page.split("\n")
        non_empty_idx = [i for i, ln in enumerate(lines) if ln.strip()]
        edge_idx = set(non_empty_idx[:3] + non_empty_idx[-3:])
        for i, ln in enumerate(lines):
            if _PAGE_NUMBER_LINE_RE.match(ln):
                continue
            if i in edge_idx and _norm(ln) in repeated:
                continue
            kept.append(ln)
        result.append("\n".join(kept))
    return result


def _unwrap_pdf_lines(text: str) -> str:
    """
    拼接 PDF 按版面宽度硬换行的句子：上一行较长（>=30 字符）且未以句末标点结尾、
    下一行不是标题/列表/表格/页码标记时，两行合并。
    """
    lines = text.split("\n")
    out: List[str] = []
    for line in lines:
        stripped = line.strip()
        if not out:
            out.append(line)
            continue
        prev = out[-1].strip()
        joinable_prev = (
            prev
            and len(prev) >= 30
            and prev[-1] not in _BOUNDARY_END_CHARS
            and not prev.startswith("|")
            and not _PAGE_MARKER_LINE_RE.match(prev)
        )
        joinable_cur = (
            stripped
            and not stripped.startswith("#")
            and not stripped.startswith("|")
            and not _PAGE_MARKER_LINE_RE.match(stripped)
            and not _BULLET_PREFIX_RE.match(stripped)
        )
        if joinable_prev and joinable_cur:
            out[-1] = out[-1].rstrip() + stripped
        else:
            out.append(line)
    return "\n".join(out)


def load_pdf_documents(file_path: Path, file_hash: str) -> List[Document]:
    """
    PDF -> 单个整篇 Document：
    逐页抽取 -> 去页眉页脚 -> 注入 [[PAGE=n]] 标记合并 -> 清洗 -> 硬换行拼接 ->
    启发式标题提升。页码标记由 semantic_splitter 消费（跨页断句修复 + chunk 级页码追溯）。
    """
    pages = _extract_pdf_pages(file_path)
    pages = _strip_repeated_page_lines(pages)

    parts: List[str] = []
    for idx, page in enumerate(pages, start=1):
        if not page.strip():
            continue
        parts.append(f"[[PAGE={idx}]]\n{page.strip()}")
    full_text = clean_text("\n".join(parts))
    full_text = _unwrap_pdf_lines(full_text)
    full_text = promote_headings(full_text)
    if not full_text:
        return []
    return [
        Document(
            page_content=full_text,
            metadata={
                "source": canonical_source_path(file_path),
                "page": 1,
                "total_pages": len(pages),
                "doc_hash": file_hash,
                "record_type": "pdf_document",
            },
        )
    ]


# =============================================================================
# DOCX 加载
# =============================================================================

_DOCX_HEADING_STYLE_RE = re.compile(r"(?:heading|标题)\s*(\d)", re.IGNORECASE)


def _docx_to_markdown(file_path: Path) -> Optional[str]:
    """
    用 python-docx 按文档顺序遍历段落与表格：
    Heading N / 标题 N 样式 -> N 级 Markdown 标题；表格 -> Markdown 表格。
    未安装 python-docx 时返回 None。
    """
    try:
        from docx import Document as DocxDocument  # type: ignore
        from docx.oxml.ns import qn  # type: ignore
        from docx.table import Table  # type: ignore
        from docx.text.paragraph import Paragraph  # type: ignore
    except ImportError:
        return None

    docx = DocxDocument(str(file_path))
    lines: List[str] = []
    for child in docx.element.body.iterchildren():
        if child.tag == qn("w:p"):
            para = Paragraph(child, docx)
            text = para.text.strip()
            if not text:
                lines.append("")
                continue
            style_name = getattr(para.style, "name", "") or ""
            match = _DOCX_HEADING_STYLE_RE.search(style_name)
            if match:
                level = min(int(match.group(1)), 6)
                lines.append(f"{'#' * level} {text}")
                lines.append("")
            else:
                lines.append(text)
        elif child.tag == qn("w:tbl"):
            table = Table(child, docx)
            rows = [
                [cell.text for cell in row.cells]
                for row in table.rows
            ]
            md = _render_table_markdown(rows)
            if md:
                lines.append("")
                lines.append(md)
                lines.append("")
    return "\n".join(lines)


def load_docx_documents(file_path: Path, file_hash: str) -> List[Document]:
    """docx -> 单个整篇 Document（优先保留标题层级/表格；退回 docx2txt 纯文本 + 启发式标题）。"""
    text = _docx_to_markdown(file_path)
    used_styles = text is not None
    if text is None:
        import docx2txt

        text = docx2txt.process(str(file_path)) or ""
    text = clean_text(text)
    if not used_styles or not _looks_like_markdown(text):
        text = promote_headings(text)
    if not text:
        return []
    return [
        Document(
            page_content=text,
            metadata={
                "source": canonical_source_path(file_path),
                "page": 0,
                "doc_hash": file_hash,
                "record_type": "docx_document",
            },
        )
    ]


# =============================================================================
# JSON 加载（问答数组逐条展开，保持既有逻辑）
# =============================================================================

def _strip_intent_suffix(text: str) -> str:
    return re.sub(r"\s*<INTENT>.*?</INTENT>\s*$", "", text, flags=re.DOTALL).strip()


def load_json_documents(file_path: Path, file_hash: str) -> List[Document]:
    """
    加载 JSON 文件。
    若是问答数组（如 instruction/input/output），则将每个问答对展开成一条独立 Document。
    """
    payload = json.loads(file_path.read_text(encoding="utf-8"))
    source = canonical_source_path(file_path)

    if isinstance(payload, list):
        documents: List[Document] = []
        for idx, item in enumerate(payload):
            if not isinstance(item, dict):
                text = clean_text(json.dumps(item, ensure_ascii=False))
                if text:
                    documents.append(
                        Document(
                            page_content=text,
                            metadata={
                                "source": source,
                                "page": idx,
                                "doc_hash": file_hash,
                                "record_type": "json_item",
                            },
                        )
                    )
                continue

            instruction = clean_text(str(item.get("instruction", "")).strip())
            user_input = clean_text(str(item.get("input", "")).strip())
            output = _strip_intent_suffix(clean_text(str(item.get("output", "")).strip()))
            system = clean_text(str(item.get("system", "")).strip())

            if instruction or output:
                parts = []
                if system:
                    parts.append(f"系统：{system}")
                if instruction:
                    parts.append(f"问题：{instruction}")
                if user_input:
                    parts.append(f"补充信息：{user_input}")
                if output:
                    parts.append(f"答案：{output}")
                page_content = "\n".join(parts).strip()
                if page_content:
                    documents.append(
                        Document(
                            page_content=page_content,
                            metadata={
                                "source": source,
                                "page": idx,
                                "doc_hash": file_hash,
                                "record_type": "qa_pair",
                                "qa_index": idx,
                            },
                        )
                    )
                continue

            text = clean_text(json.dumps(item, ensure_ascii=False))
            if text:
                documents.append(
                    Document(
                        page_content=text,
                        metadata={
                            "source": source,
                            "page": idx,
                            "doc_hash": file_hash,
                            "record_type": "json_item",
                        },
                    )
                )
        return documents

    text = clean_text(json.dumps(payload, ensure_ascii=False, indent=2))
    if not text:
        return []
    return [
        Document(
            page_content=text,
            metadata={
                "source": source,
                "page": 0,
                "doc_hash": file_hash,
                "record_type": "json_document",
            },
        )
    ]


# =============================================================================
# 统一加载入口
# =============================================================================

def _load_single_file(file_path: Path, file_hash: str) -> List[Document]:
    """按后缀分发到对应加载器，输出已清洗、带结构的 Document 列表。"""
    suffix = file_path.suffix.lower()
    if suffix == ".json":
        return load_json_documents(file_path, file_hash)
    if suffix == ".pdf":
        return load_pdf_documents(file_path, file_hash)
    if suffix == ".docx":
        return load_docx_documents(file_path, file_hash)
    if suffix in (".txt", ".md"):
        loader = TextLoader(str(file_path), encoding="utf-8", autodetect_encoding=True)
        docs = loader.load()
        out: List[Document] = []
        for doc in docs:
            text = promote_headings(clean_text(doc.page_content or ""))
            if text:
                out.append(Document(page_content=text, metadata=dict(doc.metadata)))
        return out
    if suffix == ".csv":
        docs = CSVLoader(str(file_path)).load()
        out = []
        for doc in docs:
            text = clean_text(doc.page_content or "")
            if text:
                meta = dict(doc.metadata)
                meta["record_type"] = "csv_row"
                out.append(Document(page_content=text, metadata=meta))
        return out
    return []


def load_documents(
    data_dir: str, incremental: bool = True, index_dir: Optional[str] = None
) -> Tuple[List[Document], Dict[str, int]]:
    """
    递归扫描 data_dir 下支持的文件，加载为 LangChain Document 列表。
    返回 (文档列表, 统计信息)；每条记录补充 source（相对仓库根 RAG/assets/...）/ page / doc_hash 元数据便于追溯。
    增量模式通过 index_store/doc_hash.json 中的文件哈希判定新增/变更文件。
    """
    logger.info("开始加载文档: %s", data_dir)
    root = Path(data_dir)
    if not root.exists():
        raise FileNotFoundError(f"数据目录不存在: {data_dir}")

    registry = load_hash_registry(index_dir) if incremental else {}
    records: List[Document] = []
    failed_files = 0
    skipped_unchanged = 0
    for file_path in root.rglob("*"):
        if not file_path.is_file():
            continue
        if file_path.suffix.lower() not in SUPPORTED_SUFFIXES:
            continue
        file_hash = compute_file_hash(file_path)
        source = canonical_source_path(file_path)
        tag = resolve_doc_tag(file_path, root)
        if incremental and registry.get(source) == file_hash:
            skipped_unchanged += 1
            continue
        try:
            docs = _load_single_file(file_path, file_hash)
            logger.info("成功加载文档: %s (%s 条记录)", file_path, len(docs))
        except Exception as exc:
            failed_files += 1
            logger.exception("加载失败，跳过文件: %s, err=%s", str(file_path), exc)
            continue
        for page_idx, doc in enumerate(docs):
            text = (doc.page_content or "").strip()
            if not text:
                continue
            metadata = dict(doc.metadata)
            metadata["source"] = source
            metadata["page"] = metadata.get(
                "page", metadata.get("page_number", page_idx)
            )
            metadata["doc_hash"] = metadata.get("doc_hash", file_hash)
            metadata["tag"] = tag
            records.append(Document(page_content=text, metadata=metadata))
    stats = {
        "failed_files": failed_files,
        "skipped_unchanged": skipped_unchanged,
    }
    logger.info(
        "文档加载完成: records=%s failed_files=%s skipped_unchanged=%s",
        len(records),
        failed_files,
        skipped_unchanged,
    )
    if failed_files:
        logger.warning(f"失败文件列表: {failed_files}")
    return records, stats
